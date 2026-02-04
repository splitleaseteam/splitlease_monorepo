import os
import datetime
import math
import logging
import requests
from typing import Dict, Any, Optional
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv
from flask import current_app, url_for
from werkzeug.datastructures import FileStorage
from modules.logging.logger import log_error, log_success
from modules.logging.config import LoggingConfig

# Load environment variables from backup2 root (will be mysite in PythonAnywhere)
env_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), '.env')
load_dotenv(env_path)

logger = logging.getLogger(__name__)

class CreditCardAuthGenerator:
    def __init__(self, base_dir: str = None):
        """
        Initialize the generator with paths relative to backup2 directory
        Args:
            base_dir: Base directory path, defaults to backup2 directory
        """
        if base_dir is None:
            # Default to backup2 directory (will be mysite in PythonAnywhere)
            base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))
        
        self.base_dir = base_dir
        # Use paths relative to base_dir
        self.temp_dir = os.path.join(self.base_dir, 'temp', 'credit_card_auth')
        self.template_dir = os.path.join(self.base_dir, 'templates', 'credit_card_auth')
        self.log_dir = os.path.join(self.base_dir, 'logs')
        
        # Create necessary directories with proper permissions
        self._ensure_directories()
        
        # Configure logging
        self.setup_logging()

    def _ensure_directories(self):
        """Create necessary directories with proper permissions"""
        directories = [self.temp_dir, self.template_dir, self.log_dir]
        for directory in directories:
            try:
                if not os.path.exists(directory):
                    os.makedirs(directory, mode=0o755)  # Set proper permissions
                    logger.info(f"Created directory: {directory}")
                    log_success(f"Created directory: {directory}")
            except OSError as e:
                error_msg = f"Failed to create directory {directory}: {str(e)}"
                logger.error(error_msg)
                log_error(error_msg)
                # Continue execution as directories might exist but be non-writable
                pass

    def setup_logging(self):
        """Configure logging with paths relative to base directory"""
        try:
            formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
            
            # File handler - ensure log directory exists
            log_file = os.path.join(self.log_dir, 'credit_card_auth.log')
            try:
                # Create log file if it doesn't exist
                if not os.path.exists(log_file):
                    open(log_file, 'a').close()
                    os.chmod(log_file, 0o666)  # Set proper permissions
                
                file_handler = logging.FileHandler(log_file)
                file_handler.setLevel(logging.DEBUG)
                file_handler.setFormatter(formatter)
                logger.addHandler(file_handler)
            except OSError as e:
                error_msg = f"Failed to setup file logging: {str(e)}"
                logger.error(error_msg)
                log_error(error_msg)
                # Continue with just Slack logging if file logging fails
            
            logger.setLevel(logging.DEBUG)
        except Exception as e:
            # Use print as logger might not be set up
            print(f"Failed to setup logging: {str(e)}")

    def get_date(self, date_str: str) -> Optional[datetime.datetime]:
        """Convert date string to datetime object"""
        try:
            return datetime.datetime.strptime(date_str, '%m/%d/%y')
        except ValueError:
            error_msg = f"Failed to parse date: {date_str}"
            logger.error(error_msg)
            log_error(error_msg)
            return None

    def get_formatted_date(self, date_str: str) -> str:
        """Format date string to readable format"""
        if not date_str or date_str.strip() == '':
            return ''
        date = self.get_date(date_str)
        return date.strftime('%B %d, %Y') if date else ''

    def convert_currency_to_float(self, currency_string: str) -> float:
        """Convert currency string to float"""
        try:
            cleaned_string = currency_string.replace(',', '')
            return float(cleaned_string)
        except (ValueError, TypeError):
            error_msg = f"Failed to convert currency: {currency_string}"
            logger.error(error_msg)
            log_error(error_msg)
            raise ValueError(f"Invalid currency value: {currency_string}")

    def round_down(self, value: float) -> float:
        """Round down to 2 decimal places"""
        return math.floor(value * 100) / 100

    def format_currency(self, value: float) -> str:
        """Format float as currency string"""
        return f'{value:.2f}'

    def upload_to_drive(self, file_path: str, filename: str) -> Dict[str, Any]:
        """
        Upload file to Google Drive using the Google Drive module's API
        
        Args:
            file_path: Path to the file to upload
            filename: Name of the file
            
        Returns:
            Dict containing upload result with keys:
            - success: bool
            - web_view_link: str (if successful)
            - file_id: str (if successful)
            - error: str (if failed)
        """
        try:
            # Import the upload function from the Google Drive module
            from modules.google_drive.uploader import GoogleDriveUploader
            
            # Initialize uploader with the same base directory
            uploader = GoogleDriveUploader(self.base_dir)
            
            # Upload directly using the uploader
            with open(file_path, 'rb') as file:
                result = uploader.upload_file(
                    file=file,
                    file_name=filename,
                    mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
            
            if result.get('success'):
                success_msg = f"Successfully uploaded file to Google Drive: {filename}"
                log_success(success_msg)
                return {
                    'success': True,
                    'web_view_link': result.get('web_view_link'),
                    'file_id': result.get('file_id')
                }
            else:
                error_msg = f"Failed to upload to Google Drive: {result.get('error')}"
                logger.error(error_msg)
                log_error(error_msg)
                return {
                    'success': False,
                    'error': error_msg
                }
                
        except Exception as e:
            error_msg = f"Error uploading to Google Drive: {str(e)}"
            logger.error(error_msg)
            log_error(error_msg)
            return {
                'success': False,
                'error': error_msg
            }

    def create_authorization(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Create a new recurring credit card authorization form"""
        try:
            logger.info(f"Creating authorization for agreement {data.get('Agreement Number', 'UNKNOWN')}")
            log_success(f"Creating authorization for agreement {data.get('Agreement Number', 'UNKNOWN')}")
            
            # Use absolute path for template
            template_path = os.path.join(self.template_dir, 'recurringcreditcardauthorizationprorated.docx')
            if not os.path.exists(template_path):
                error_msg = f"Template file not found at {template_path}"
                logger.error(error_msg)
                log_error(error_msg)
                return {'success': False, 'error': error_msg}

            # Create initial document structure
            try:
                document = Document(template_path)
                docum = document.add_paragraph()
                docu = docum.add_run()
            except Exception as e:
                error_msg = f"Failed to create initial document: {str(e)}"
                logger.error(error_msg)
                log_error(error_msg)
                return {'success': False, 'error': f"Document creation error: {str(e)}"}

            # Initialize template
            try:
                doc = DocxTemplate(template_path)
            except Exception as e:
                error_msg = f"Failed to initialize template: {str(e)}"
                logger.error(error_msg)
                log_error(error_msg)
                return {'success': False, 'error': f"Template initialization error: {str(e)}"}

            # Process currency values
            try:
                four_week_rent = self.round_down(self.convert_currency_to_float(data['Four Week Rent']))
                maintenance_fee = self.round_down(self.convert_currency_to_float(data['Maintenance Fee']))
                damage_deposit = self.round_down(self.convert_currency_to_float(data['Damage Deposit']))
                splitlease_credit = self.round_down(self.convert_currency_to_float(data['Splitlease Credit']))
                last_payment_rent = self.round_down(self.convert_currency_to_float(data['Last Payment Rent']))

                logger.debug(f"Processed currency values - Four Week Rent: {four_week_rent}, "
                           f"Maintenance Fee: {maintenance_fee}, Damage Deposit: {damage_deposit}")

                # Calculate totals
                total_first_payment = self.round_down(four_week_rent + maintenance_fee + damage_deposit)
                total_second_payment = self.round_down(four_week_rent + maintenance_fee)
                total_last_payment = self.round_down(last_payment_rent + maintenance_fee - splitlease_credit)

                logger.debug(f"Calculated payment totals - First: {total_first_payment}, "
                           f"Second: {total_second_payment}, Last: {total_last_payment}")

            except ValueError as e:
                error_msg = f"Error processing currency values: {str(e)}"
                logger.error(error_msg)
                log_error(error_msg)
                return {'success': False, 'error': f"Currency processing error: {str(e)}"}

            # Prepare template data
            items = {
                'agreement_number': data['Agreement Number'],
                'host_name': data['Host Name'],
                'guest_name': data['Guest Name'],
                'maintenancefee': self.format_currency(maintenance_fee),
                'weeks_number': data['Weeks Number'],
                'ListingDescription': data['Listing Description'],
                'fourweekrent': self.format_currency(four_week_rent),
                'damagedeposit': self.format_currency(damage_deposit),
                'totalfirstpayment': self.format_currency(total_first_payment),
                'penultimateweeknumber': data['Penultimate Week Number'],
                'totalsecondpayment': self.format_currency(total_second_payment),
                'slcredit': self.format_currency(splitlease_credit),
                'lastpaymenttotal': self.format_currency(total_last_payment),
                'numberofpayments': data['Number of Payments'],
                'lastpaymentweeks': data['Last Payment Weeks'],
                'lastpaymentrent': self.format_currency(last_payment_rent)
            }

            # Render and save document
            try:
                doc.render(items)
                filename = f'recurring_credit_card_auth-prorated-{data["Agreement Number"]}.docx'
                file_path = os.path.join(self.temp_dir, filename)
                doc.save(file_path)
                
                # Set proper permissions for the generated file
                os.chmod(file_path, 0o666)
                
                # Upload to Google Drive using the module's API
                upload_result = self.upload_to_drive(file_path, filename)
                
                if not upload_result['success']:
                    return {
                        'success': True,
                        'file_path': file_path,
                        'filename': filename,
                        'drive_upload_error': upload_result.get('error')
                    }
                
                success_msg = f"Successfully created and uploaded authorization document: {filename}"
                logger.info(success_msg)
                log_success(success_msg)
                return {
                    'success': True,
                    'file_path': file_path,
                    'filename': filename,
                    'web_view_link': upload_result.get('web_view_link'),
                    'file_id': upload_result.get('file_id')
                }
            except Exception as e:
                error_msg = f"Failed to save document: {str(e)}"
                logger.error(error_msg)
                log_error(error_msg)
                return {'success': False, 'error': f"Document save error: {str(e)}"}

        except Exception as e:
            error_msg = f"Error creating authorization: {str(e)}"
            logger.error(error_msg)
            log_error(error_msg)
            return {
                'success': False,
                'error': error_msg
            }

class CreditCardAuthNonProratedGenerator:
    def __init__(self, base_dir: str = None):
        if base_dir is None:
            base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))

        self.base_dir = base_dir
        self.temp_dir = os.path.join(self.base_dir, 'temp', 'credit_card_auth_nonprorated')
        self.template_dir = os.path.join(self.base_dir, 'templates', 'credit_card_auth_nonprorated')
        self.log_dir = os.path.join(self.base_dir, 'logs')

        self._ensure_directories()
        self.setup_logging()

    def _ensure_directories(self):
        directories = [self.temp_dir, self.template_dir, self.log_dir]
        for directory in directories:
            try:
                if not os.path.exists(directory):
                    os.makedirs(directory, mode=0o755)
                    logger.info(f"Created directory: {directory}")
                    log_success(f"Created directory: {directory}")
            except OSError as e:
                error_msg = f"Failed to create directory {directory}: {str(e)}"
                logger.error(error_msg)
                log_error(error_msg)
                pass

    def setup_logging(self):
        try:
            formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')

            log_file = os.path.join(self.log_dir, 'credit_card_auth_nonprorated.log')
            try:
                if not os.path.exists(log_file):
                    open(log_file, 'a').close()
                    os.chmod(log_file, 0o666)

                file_handler = logging.FileHandler(log_file)
                file_handler.setLevel(logging.DEBUG)
                file_handler.setFormatter(formatter)
                logger.addHandler(file_handler)
            except OSError as e:
                error_msg = f"Failed to setup file logging: {str(e)}"
                logger.error(error_msg)
                log_error(error_msg)

            logger.setLevel(logging.DEBUG)
        except Exception as e:
            print(f"Failed to setup logging: {str(e)}")

    def convert_currency_to_float(self, currency_string: str) -> float:
        try:
            cleaned_string = currency_string.replace(',', '')
            return float(cleaned_string)
        except (ValueError, TypeError):
            error_msg = f"Failed to convert currency: {currency_string}"
            logger.error(error_msg)
            log_error(error_msg)
            raise ValueError(f"Invalid currency value: {currency_string}")

    def round_down(self, value: float) -> float:
        return math.floor(value * 100) / 100

    def format_currency(self, value: float) -> str:
        return f'{value:.2f}'

    def calculate_payments(self, data: Dict[str, Any]) -> Dict[str, float]:
        try:
            four_week_rent = self.round_down(self.convert_currency_to_float(data['Four Week Rent']))
            maintenance_fee = self.round_down(self.convert_currency_to_float(data['Maintenance Fee']))
            damage_deposit = self.round_down(self.convert_currency_to_float(data['Damage Deposit']))
            splitlease_credit = self.round_down(self.convert_currency_to_float(data['Splitlease Credit']))
            last_payment_rent = self.round_down(self.convert_currency_to_float(data['Last Payment Rent']))

            total_first_payment = self.round_down(four_week_rent + maintenance_fee + damage_deposit)
            total_second_payment = self.round_down(four_week_rent + maintenance_fee)
            total_last_payment = self.round_down(last_payment_rent + maintenance_fee - splitlease_credit)

            return {
                'four_week_rent': four_week_rent,
                'maintenance_fee': maintenance_fee,
                'damage_deposit': damage_deposit,
                'splitlease_credit': splitlease_credit,
                'last_payment_rent': last_payment_rent,
                'total_first_payment': total_first_payment,
                'total_second_payment': total_second_payment,
                'total_last_payment': total_last_payment
            }
        except Exception as e:
            error_msg = f"Error calculating payments: {str(e)}"
            logger.error(error_msg)
            log_error(error_msg)
            raise

    def upload_to_drive(self, file_path: str, filename: str) -> Dict[str, Any]:
        try:
            from modules.google_drive.uploader import GoogleDriveUploader

            uploader = GoogleDriveUploader(self.base_dir)

            with open(file_path, 'rb') as file:
                result = uploader.upload_file(
                    file=file,
                    file_name=filename,
                    mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

            if result.get('success'):
                success_msg = f"Successfully uploaded file to Google Drive: {filename}"
                log_success(success_msg)
                return {
                    'success': True,
                    'web_view_link': result.get('web_view_link'),
                    'file_id': result.get('file_id')
                }
            else:
                error_msg = f"Failed to upload to Google Drive: {result.get('error')}"
                logger.error(error_msg)
                log_error(error_msg)
                return {
                    'success': False,
                    'error': error_msg
                }

        except Exception as e:
            error_msg = f"Error uploading to Google Drive: {str(e)}"
            logger.error(error_msg)
            log_error(error_msg)
            return {
                'success': False,
                'error': error_msg
            }

    def create_authorization(self, data: Dict[str, Any]) -> Dict[str, Any]:
        try:
            logger.info(f"Creating authorization for agreement {data.get('Agreement Number', 'UNKNOWN')}")
            log_success(f"Creating authorization for agreement {data.get('Agreement Number', 'UNKNOWN')}")

            template_path = os.path.join(self.template_dir, 'recurringcreditcardauthorization.docx')
            if not os.path.exists(template_path):
                error_msg = f"Template file not found at {template_path}"
                logger.error(error_msg)
                log_error(error_msg)
                return {'success': False, 'error': error_msg}

            try:
                document = Document(template_path)
                docum = document.add_paragraph()
                docu = docum.add_run()
            except Exception as e:
                error_msg = f"Failed to create initial document: {str(e)}"
                logger.error(error_msg)
                log_error(error_msg)
                return {'success': False, 'error': f"Document creation error: {str(e)}"}

            try:
                doc = DocxTemplate(template_path)
            except Exception as e:
                error_msg = f"Failed to initialize template: {str(e)}"
                logger.error(error_msg)
                log_error(error_msg)
                return {'success': False, 'error': f"Template initialization error: {str(e)}"}

            try:
                payments = self.calculate_payments(data)
            except Exception as e:
                error_msg = f"Error calculating payments: {str(e)}"
                logger.error(error_msg)
                log_error(error_msg)
                return {'success': False, 'error': f"Payment calculation error: {str(e)}"}

            items = {
                'agreement_number': data['Agreement Number'],
                'host_name': data['Host Name'],
                'guest_name': data['Guest Name'],
                'maintenancefee': self.format_currency(payments['maintenance_fee']),
                'weeks_number': data['Weeks Number'],
                'ListingDescription': data['Listing Description'],
                'fourweekrent': self.format_currency(payments['four_week_rent']),
                'damagedeposit': self.format_currency(payments['damage_deposit']),
                'totalfirstpayment': self.format_currency(payments['total_first_payment']),
                'penultimateweeknumber': data['Penultimate Week Number'],
                'totalsecondpayment': self.format_currency(payments['total_second_payment']),
                'slcredit': self.format_currency(payments['splitlease_credit']),
                'lastpaymenttotal': self.format_currency(payments['total_last_payment']),
                'numberofpayments': data['Number of Payments'],
                'lastpaymentweeks': data['Last Payment Weeks'],
                'lastpaymentrent': self.format_currency(payments['last_payment_rent'])
            }

            try:
                doc.render(items)
                filename = f'recurring_credit_card_auth-nonprorated-{data["Agreement Number"]}.docx'
                file_path = os.path.join(self.temp_dir, filename)
                doc.save(file_path)
                os.chmod(file_path, 0o666)

                upload_result = self.upload_to_drive(file_path, filename)

                if not upload_result['success']:
                    return {
                        'success': True,
                        'file_path': file_path,
                        'filename': filename,
                        'drive_upload_error': upload_result.get('error')
                    }

                success_msg = f"Successfully created and uploaded authorization document: {filename}"
                logger.info(success_msg)
                log_success(success_msg)
                return {
                    'success': True,
                    'file_path': file_path,
                    'filename': filename,
                    'web_view_link': upload_result.get('web_view_link'),
                    'file_id': upload_result.get('file_id')
                }
            except Exception as e:
                error_msg = f"Failed to save document: {str(e)}"
                logger.error(error_msg)
                log_error(error_msg)
                return {'success': False, 'error': f"Document save error: {str(e)}"}

        except Exception as e:
            error_msg = f"Error creating authorization: {str(e)}"
            logger.error(error_msg)
            log_error(error_msg)
            return {
                'success': False,
                'error': error_msg
            }
