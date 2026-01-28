import os
import datetime
import io
import logging
from typing import Dict, Any, Optional
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
import requests
from modules.google_drive.uploader import GoogleDriveUploader
from io import BytesIO
from modules.logging.error_logger import log_error
from modules.logging.success_logger import log_success

logger = logging.getLogger(__name__)

class HostPayoutGenerator:
    def __init__(self, base_dir: str):
        """Initialize the host payout generator.
            
        Args:
            base_dir: Base directory for the project (project root)
        """
        self.base_dir = base_dir
        self.temp_dir = os.path.join(base_dir, 'temp', 'host_payout')
        self.template_dir = os.path.join(base_dir, 'templates', 'host_payout')
        self.log_dir = os.path.join(base_dir, 'logs')
        
        # Create necessary directories
        os.makedirs(self.temp_dir, exist_ok=True)
        os.makedirs(self.template_dir, exist_ok=True)
        os.makedirs(self.log_dir, exist_ok=True)
        
        # Configure logging
        self.setup_logging()
        
        # Initialize Google Drive uploader
        self.drive_uploader = GoogleDriveUploader(base_dir)

    def setup_logging(self):
        """Configure logging to both file and Slack"""
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        
        # File handler
        file_handler = logging.FileHandler(os.path.join(self.log_dir, 'host_payout.log'))
        file_handler.setLevel(logging.DEBUG)
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)
        
        logger.setLevel(logging.DEBUG)

    def get_date(self, date_str: str) -> Optional[datetime.datetime]:
        """Convert date string to datetime object"""
        try:
            # Try MM/DD/YY format first
            return datetime.datetime.strptime(date_str, '%m/%d/%y')
        except ValueError:
            try:
                # Try YYYY-MM-DD format
                return datetime.datetime.strptime(date_str, '%Y-%m-%d')
            except ValueError as e:
                logger.error(f"Date parsing error for {date_str}: {str(e)}")
                return None

    def get_formatted_date(self, date_str: str) -> str:
        """Format date string to readable format"""
        if not date_str or date_str.strip() == '':
            return ''
        date = self.get_date(date_str)
        return date.strftime('%B %d, %Y') if date else ''

    def format_currency(self, value: str) -> str:
        """Format value as currency with thousands separator"""
        if not value or value.strip() == '':
            return ''  # Return empty string without warning
        try:
            # Remove any currency symbols and whitespace
            clean_value = value.replace('$', '').replace(',', '').strip()
            # Format with comma as thousands separator
            return '${:,.2f}'.format(float(clean_value))
        except (ValueError, TypeError) as e:
            logger.warning(f"Currency formatting error for {value}: {str(e)}")
            return value

    def remove_empty_rows(self, table) -> None:
        """Remove empty rows from the payout schedule table"""
        try:
            rows_to_delete = []
            for i, row in enumerate(table.rows[1:], start=1):  # Skip header row
                if all(cell.text.strip() == '' for cell in row.cells[1:4]):  # Check Rent#, Total#, Maintenance Fee#
                    rows_to_delete.append(i)

            for row_index in reversed(rows_to_delete):
                table._tbl.remove(table.rows[row_index]._tr)
        except Exception as e:
            logger.error(f"Error removing empty rows: {str(e)}")

    def create_schedule(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Create a new host payout schedule"""
        try:
            template_path = os.path.join(self.template_dir, 'hostpayoutscheduleform.docx')
            if not os.path.exists(template_path):
                error_msg = f"Template file not found at {template_path}"
                logger.error(error_msg)
                log_error(error_msg)
                return {'success': False, 'error': error_msg, 'returned_error': 'yes'}

            doc = DocxTemplate(template_path)
            
            # Prepare base template data
            template_data = {
                "address": data.get("Address", ""),
                "agreement_number": data.get("Agreement Number", ""),
                "host_email": data.get("Host Email", ""),
                "host_name": data.get("Host Name", ""),
                "host_phone": data.get("Host Phone", ""),
                "payout_number": data.get("Payout Number", ""),
            }

            # Add payment schedule data
            for i in range(1, 14):
                template_data.update({
                    f"date{i}": self.get_formatted_date(data.get(f'Date{i}', '')),
                    f"rent{i}": self.format_currency(data.get(f"Rent{i}", "")),
                    f"total{i}": self.format_currency(data.get(f"Total{i}", "")),
                    f"maintenance_fee{i}": self.format_currency(data["Maintenance Fee"]) 
                        if data.get(f"Rent{i}", "").strip() else ""
                })

            # Render template
            doc.render(template_data)

            # Process the rendered document
            temp_buffer = BytesIO()
            doc.save(temp_buffer)
            temp_buffer.seek(0)

            # Create final document
            document = Document(temp_buffer)
            if document.tables:
                self.remove_empty_rows(document.tables[0])

            # Save final document
            agreement_number = data.get('Agreement Number', 'unknown')
            filename = f'host_payout_schedule-{agreement_number}.docx'
            file_path = os.path.join(self.temp_dir, filename)
            document.save(file_path)

            # Upload to Google Drive
            try:
                mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'  # MIME type for .docx
                upload_result = self.drive_uploader.upload_file(file_path, filename, mime_type)
                if upload_result.get('success'):
                    success_msg = f"Successfully created and uploaded host payout agreement: {filename}"
                    logger.info(success_msg)
                    log_success(success_msg)
                    
                    return {
                        'success': True,
                        'file_path': file_path,
                        'filename': filename,
                        'file_id': upload_result.get('file_id'),
                        'web_view_link': upload_result.get('web_view_link'),
                        'returned_error': 'no'
                    }
                else:
                    error_msg = f"Created file but failed to upload to Google Drive: {upload_result.get('error')}"
                    logger.error(error_msg)
                    log_error(error_msg)
            except Exception as upload_error:
                error_msg = f"Error uploading to Google Drive: {str(upload_error)}"
                logger.error(error_msg)
                log_error(error_msg)

            # Return just the local file if Google Drive upload failed
            return {
                'success': True,
                'file_path': file_path,
                'filename': filename,
                'returned_error': 'no'
            }
        except Exception as e:
            error_msg = f"Error creating host payout schedule: {str(e)}"
            logger.error(error_msg)
            log_error(error_msg)
            return {
                'success': False,
                'error': error_msg,
                'returned_error': 'yes'
            }
