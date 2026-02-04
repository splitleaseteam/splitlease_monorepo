import os
import datetime
import math
import logging
import requests
from typing import Dict, Any, Optional
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv
from flask import current_app, url_for
from werkzeug.datastructures import FileStorage
from modules.logging.logger import log_error, log_success
from modules.logging.config import LoggingConfig

# Load environment variables from backup2 root
env_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), '.env')
load_dotenv(env_path)

logger = logging.getLogger(__name__)

class CreditCardAuthNonProratedGenerator:
    def __init__(self, base_dir: str = None):
        if base_dir is None:
            base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))

        self.base_dir = base_dir
        self.temp_dir = os.path.join(self.base_dir, 'temp', 'credit_card_auth_nonprorated')
        self.template_dir = os.path.join(self.base_dir, 'templates', 'credit_card_auth_nonprorated')
        self.log_dir = os.path.join(self.base_dir, 'logs')

        self._ensure_directories()
        self.setup_logging()

    def _ensure_directories(self):
        directories = [self.temp_dir, self.template_dir, self.log_dir]
        for directory in directories:
            try:
                if not os.path.exists(directory):
                    os.makedirs(directory, mode=0o755)
                    log_success(f"Created directory: {directory}")
            except OSError as e:
                log_error(f"Failed to create directory {directory}: {str(e)}")
                pass

    def setup_logging(self):
        try:
            formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')

            log_file = os.path.join(self.log_dir, 'credit_card_auth_nonprorated.log')
            try:
                if not os.path.exists(log_file):
                    open(log_file, 'a').close()
                    os.chmod(log_file, 0o666)

                file_handler = logging.FileHandler(log_file)
                file_handler.setLevel(logging.DEBUG)
                file_handler.setFormatter(formatter)
                logger.addHandler(file_handler)
            except OSError as e:
                log_error(f"Failed to setup file logging: {str(e)}")

            # Set up console handler for local debugging
            console_handler = logging.StreamHandler()
            console_handler.setLevel(logging.DEBUG)
            console_handler.setFormatter(formatter)
            logger.addHandler(console_handler)

            logger.setLevel(logging.DEBUG)
            
            # Test the logging setup
            log_success("Credit Card Auth NonProrated logging initialized")
        except Exception as e:
            print(f"Failed to setup logging: {str(e)}")

    def convert_currency_to_float(self, currency_string: str) -> float:
        try:
            cleaned_string = currency_string.replace(',', '')
            return float(cleaned_string)
        except (ValueError, TypeError):
            log_error(f"Failed to convert currency: {currency_string}")
            raise ValueError(f"Invalid currency value: {currency_string}")

    def round_down(self, value: float) -> float:
        return math.floor(value * 100) / 100

    def format_currency(self, value: float) -> str:
        return f'{value:.2f}'

    def calculate_payments(self, data: Dict[str, Any]) -> Dict[str, float]:
        try:
            four_week_rent = self.round_down(self.convert_currency_to_float(data['Four Week Rent']))
            maintenance_fee = self.round_down(self.convert_currency_to_float(data['Maintenance Fee']))
            damage_deposit = self.round_down(self.convert_currency_to_float(data['Damage Deposit']))
            splitlease_credit = self.round_down(self.convert_currency_to_float(data['Splitlease Credit']))
            last_payment_rent = self.round_down(self.convert_currency_to_float(data['Last Payment Rent']))

            total_first_payment = self.round_down(four_week_rent + maintenance_fee + damage_deposit)
            total_second_payment = self.round_down(four_week_rent + maintenance_fee)
            total_last_payment = self.round_down(last_payment_rent + maintenance_fee - splitlease_credit)

            return {
                'four_week_rent': four_week_rent,
                'maintenance_fee': maintenance_fee,
                'damage_deposit': damage_deposit,
                'splitlease_credit': splitlease_credit,
                'last_payment_rent': last_payment_rent,
                'total_first_payment': total_first_payment,
                'total_second_payment': total_second_payment,
                'total_last_payment': total_last_payment
            }
        except Exception as e:
            log_error(f"Error calculating payments: {str(e)}")
            raise

    def upload_to_drive(self, file_path: str, filename: str) -> Dict[str, Any]:
        try:
            from modules.google_drive.uploader import GoogleDriveUploader

            uploader = GoogleDriveUploader(self.base_dir)

            with open(file_path, 'rb') as file:
                result = uploader.upload_file(
                    file=file,
                    file_name=filename,
                    mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

            if result.get('success'):
                log_success(f"Successfully uploaded {filename} to Google Drive")
                return {
                    'success': True,
                    'web_view_link': result.get('web_view_link'),
                    'file_id': result.get('file_id')
                }
            else:
                error_msg = f"Failed to upload to Google Drive: {result.get('error')}"
                log_error(error_msg)
                return {
                    'success': False,
                    'error': error_msg
                }

        except Exception as e:
            error_msg = f"Error uploading to Google Drive: {str(e)}"
            log_error(error_msg)
            return {
                'success': False,
                'error': error_msg
            }

    def create_authorization(self, data: Dict[str, Any]) -> Dict[str, Any]:
        try:
            log_success(f"Creating authorization for agreement {data.get('Agreement Number', 'UNKNOWN')}")

            template_path = os.path.join(self.template_dir, 'recurringcreditcardauthorization.docx')
            if not os.path.exists(template_path):
                error_msg = f"Template file not found at {template_path}"
                log_error(error_msg)
                return {'success': False, 'error': error_msg}

            try:
                document = Document(template_path)
                docum = document.add_paragraph()
                docu = docum.add_run()
            except Exception as e:
                log_error(f"Failed to create initial document: {str(e)}")
                return {'success': False, 'error': f"Document creation error: {str(e)}"}

            try:
                doc = DocxTemplate(template_path)
            except Exception as e:
                log_error(f"Failed to initialize template: {str(e)}")
                return {'success': False, 'error': f"Template initialization error: {str(e)}"}

            try:
                payments = self.calculate_payments(data)
            except Exception as e:
                log_error(f"Error calculating payments: {str(e)}")
                return {'success': False, 'error': f"Payment calculation error: {str(e)}"}

            items = {
                'agreement_number': data['Agreement Number'],
                'host_name': data['Host Name'],
                'guest_name': data['Guest Name'],
                'maintenancefee': self.format_currency(payments['maintenance_fee']),
                'weeks_number': data['Weeks Number'],
                'ListingDescription': data['Listing Description'],
                'fourweekrent': self.format_currency(payments['four_week_rent']),
                'damagedeposit': self.format_currency(payments['damage_deposit']),
                'totalfirstpayment': self.format_currency(payments['total_first_payment']),
                'penultimateweeknumber': data['Penultimate Week Number'],
                'totalsecondpayment': self.format_currency(payments['total_second_payment']),
                'slcredit': self.format_currency(payments['splitlease_credit']),
                'lastpaymenttotal': self.format_currency(payments['total_last_payment']),
                'numberofpayments': data['Number of Payments'],
                'lastpaymentweeks': data['Last Payment Weeks'],
                'lastpaymentrent': self.format_currency(payments['last_payment_rent'])
            }

            try:
                doc.render(items)
                filename = f'recurring_credit_card_auth-nonprorated-{data["Agreement Number"]}.docx'
                file_path = os.path.join(self.temp_dir, filename)
                doc.save(file_path)
                os.chmod(file_path, 0o666)

                upload_result = self.upload_to_drive(file_path, filename)

                if not upload_result['success']:
                    return {
                        'success': True,
                        'file_path': file_path,
                        'filename': filename,
                        'drive_upload_error': upload_result.get('error')
                    }

                log_success(f"Successfully created and uploaded authorization document: {filename}")
                return {
                    'success': True,
                    'file_path': file_path,
                    'filename': filename,
                    'web_view_link': upload_result.get('web_view_link'),
                    'file_id': upload_result.get('file_id')
                }
            except Exception as e:
                log_error(f"Failed to save document: {str(e)}")
                return {'success': False, 'error': f"Document save error: {str(e)}"}

        except Exception as e:
            error_msg = f"Error creating authorization: {str(e)}"
            log_error(error_msg)
            return {
                'success': False,
                'error': error_msg
            }
