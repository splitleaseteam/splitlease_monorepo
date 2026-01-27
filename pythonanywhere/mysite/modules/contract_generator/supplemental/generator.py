import os
import datetime
import io
import logging
from typing import Dict, Any, Optional
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
import requests
from modules.google_drive.uploader import GoogleDriveUploader
from modules.logging.error_logger import log_error
from modules.logging.success_logger import log_success

logger = logging.getLogger(__name__)

# Webhook URLs handled internally by logging helpers

class SupplementalAgreementGenerator:
    def __init__(self, base_dir: str):
        """Initialize the supplemental agreement generator.
        
        Args:
            base_dir: Base directory for all files
        """
        self.base_dir = base_dir
        self.module_name = 'supplemental'
        self.temp_dir = os.path.join(base_dir, 'temp', self.module_name)
        self.template_dir = os.path.join(base_dir, 'modules', 'templates', self.module_name)
        self.log_dir = os.path.join(base_dir, 'logs')
        
        # Create necessary directories
        os.makedirs(self.temp_dir, exist_ok=True)
        os.makedirs(self.template_dir, exist_ok=True)
        os.makedirs(self.log_dir, exist_ok=True)
        
        # Configure logging
        self.setup_logging()
        
        # Initialize Google Drive uploader
        self.drive_uploader = GoogleDriveUploader(base_dir)
        
        # Configure specific paths for supplemental agreements
        self.template_path = os.path.join(self.template_dir, 'supplementalagreement.docx')
        
        # Verify template exists
        if not os.path.exists(self.template_path):
            error_msg = f"Template file not found at {self.template_path}"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)

    def setup_logging(self):
        """Configure logging to both file and Slack"""
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        
        # File handler
        file_handler = logging.FileHandler(os.path.join(self.log_dir, f'{self.module_name}.log'))
        file_handler.setLevel(logging.DEBUG)
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)
        
        logger.setLevel(logging.DEBUG)

    def get_date(self, date_str: str) -> datetime.datetime:
        """Convert date string to datetime object"""
        if not date_str or date_str.strip() == '':
            logger.warning("Empty date string provided")
            return None
            
        try:
            # Try MM/DD/YY format first
            logger.debug(f"Attempting to parse date {date_str} in MM/DD/YY format")
            return datetime.datetime.strptime(date_str, '%m/%d/%y')
        except ValueError:
            try:
                # Try YYYY-MM-DD format
                logger.debug(f"Attempting to parse date {date_str} in YYYY-MM-DD format")
                return datetime.datetime.strptime(date_str, '%Y-%m-%d')
            except ValueError as e:
                logger.error(f"Date parsing error for {date_str}: {str(e)}")
                return None

    def get_formatted_date(self, date_str: str) -> str:
        """Format date string to readable format"""
        if not date_str or date_str.strip() == '':
            logger.warning("Empty date string provided")
            return ''
            
        date = self.get_date(date_str)
        if date is None:
            logger.error(f"Could not parse date: {date_str}")
            return date_str
            
        formatted_date = date.strftime('%B %d, %Y')
        logger.debug(f"Formatted date {date_str} to {formatted_date}")
        return formatted_date

    def process_images(self, doc: DocxTemplate, data: Dict[str, Any]) -> Dict[str, Any]:
        """Process and prepare images for the document"""
        logger.info("Starting image processing")
        images = {}
        
        for i in range(1, 4):
            image_key = f'image{i}'
            if image_key in data and data[image_key]:
                try:
                    logger.debug(f"Processing {image_key}")
                    image_url = data[image_key]
                    
                    try:
                        # Download image from URL
                        response = requests.get(image_url)
                        response.raise_for_status()
                        
                        # Create BytesIO object from the response content
                        image_stream = io.BytesIO(response.content)
                        
                        # Create InlineImage with the stream
                        images[image_key] = InlineImage(doc, image_stream, width=Mm(48))
                        logger.info(f"Successfully processed {image_key}")
                        
                    except Exception as download_error:
                        logger.error(f"Error downloading image for {image_key}: {str(download_error)}")
                        logger.exception(download_error)
                        images[image_key] = ''
                        
                except Exception as e:
                    logger.error(f"Error in image processing for {image_key}: {str(e)}")
                    logger.exception(e)
                    images[image_key] = ''
            else:
                images[image_key] = ''
        
        logger.info(f"Completed image processing. Processed {len(images)} images")
        return images

    def upload_to_drive(self, file_path: str, file_name: Optional[str] = None) -> Optional[Dict[str, Any]]:
        """Upload file to Google Drive"""
        try:
            uploader = self.drive_uploader
            if not uploader:
                error_msg = f"Failed to upload {file_name or file_path}: Google Drive uploader not initialized"
                logger.error(error_msg)
                log_error(error_msg)
                return None

            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            result = uploader.upload_file(file_path, file_name or os.path.basename(file_path), mime_type)
            
            if result.get('success'):
                success_msg = f"Successfully uploaded {file_name or file_path} to Google Drive"
                logger.info(f"Successfully uploaded file to Google Drive: {result.get('file_id')}")
                log_success(success_msg)
                return result
            else:
                error_msg = f"Failed to upload {file_name or file_path} to Google Drive: {result.get('error')}"
                logger.error(f"Failed to upload file to Google Drive: {result.get('error')}")
                log_error(error_msg)
                return None

        except Exception as e:
            error_msg = f"Error uploading {file_name or file_path} to Google Drive: {str(e)}"
            logger.error(f"Error uploading to Google Drive: {str(e)}")
            log_error(error_msg)
            return None

    def remove_empty_pages(self, file_path: str) -> None:
        """Remove trailing empty paragraphs and adjust sections to prevent empty pages"""
        from docx.enum.section import WD_SECTION
        from docx.oxml.shared import qn
        from docx import Document
        
        try:
            doc = Document(file_path)
            
            # Remove trailing empty paragraphs
            while doc.paragraphs and not doc.paragraphs[-1].text.strip():
                last_para = doc.paragraphs[-1]
                last_para._element.getparent().remove(last_para._element)
            
            # Check section breaks
            if len(doc.sections) > 1:
                # Remove empty last section if it's continuous
                last_section = doc.sections[-1]
                sectPr = last_section._sectPr
                type_attr = sectPr.xpath('.//w:type/@w:val')
                if type_attr and type_attr[0] == 'continuous':
                    doc.delete_section(-1)
            
            # Remove empty paragraph after tables/images
            for element in reversed(doc.element.body):
                if element.tag.endswith('sectPr'):
                    continue
                if len(element) == 0 or not element.text.strip():
                    doc.element.body.remove(element)
                else:
                    break
            
            doc.save(file_path)
            logger.info("Successfully cleaned up document formatting")
            
        except Exception as e:
            logger.error(f"Error during post-processing: {str(e)}")
            raise

    def create_agreement(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Create a new supplemental agreement and upload to Google Drive"""
        try:
            # Log input data
            logger.info(f"Creating supplemental agreement for Agreement Number: {data.get('Agreement Number', 'unknown')}")
            
            # Load template
            logger.info("Loading template file")
            try:
                doc = DocxTemplate(self.template_path)
                logger.info("Template loaded successfully")
            except Exception as e:
                error_msg = f"Failed to load template: {str(e)}"
                logger.error(error_msg)
                logger.exception(e)
                log_error(error_msg)
                return {'success': False, 'error': error_msg}
            
            # Process images
            images = self.process_images(doc, data)
            
            # Prepare template data
            logger.info("Preparing template data")
            template_data = {
                'agreement_number': str(data.get('Agreement Number', '')),
                'start_date': self.get_formatted_date(data.get('Check in Date', '')),
                'end_date': self.get_formatted_date(data.get('Check Out Date', '')),
                'weeks_number': str(data.get('Number of weeks', '')),
                'guest_allowed': str(data.get('Guests Allowed', '')),
                'guests_allowed': str(data.get('Guests Allowed', '')),
                'host_name': str(data.get('Host Name', '')),
                'listing_description': str(data.get('Listing Description', '')),
                'listing_title': str(data.get('Listing Title', '')),
                'spacedetails': str(data.get('Space Details', '')),
                'location': str(data.get('Location', '')),
                'type_of_space': str(data.get('Type of Space', '')),
                'supplement_number': str(data.get('Supplemental Number', '')),
                **images
            }
            
            # Log template data
            for key, value in template_data.items():
                if isinstance(value, InlineImage):
                    logger.debug(f"  {key}: <InlineImage>")
                else:
                    logger.debug(f"  {key}: {value}")

            # Render document
            logger.info("Rendering template")
            try:
                doc.render(template_data)
                logger.info("Template rendered successfully")
            except Exception as render_error:
                error_msg = f"Template render error: {str(render_error)}"
                logger.error(error_msg)
                logger.exception(render_error)
                log_error(error_msg)
                return {'success': False, 'error': error_msg}

            # Save and upload document
            agreement_number = data.get('Agreement Number', 'unknown')
            filename = f'supplement_agreement-{agreement_number}.docx'
            file_path = os.path.join(self.temp_dir, filename)
            
            logger.info(f"Saving document to {file_path}")
            try:
                doc.save(file_path)
                logger.info("Document saved successfully")
                
                # Add post-processing to remove empty pages
                try:
                    self.remove_empty_pages(file_path)
                except Exception as post_process_error:
                    error_msg = f"Post-processing error: {str(post_process_error)}"
                    logger.warning(error_msg)
                    logger.exception(post_process_error)
                
                # Upload to Google Drive
                logger.info("Uploading to Google Drive")
                try:
                    drive_result = self.upload_to_drive(file_path, filename)
                    
                    if not drive_result:
                        logger.error("Failed to upload to Google Drive")
                        return {
                            'success': True,
                            'file_path': file_path,
                            'filename': filename,
                            'drive_error': 'Failed to upload to Google Drive'
                        }
                    
                    logger.info("Successfully uploaded to Google Drive")
                    return {
                        'success': True,
                        'file_path': file_path,
                        'filename': filename,
                        'web_view_link': drive_result['web_view_link']
                    }
                    
                except Exception as upload_error:
                    error_msg = f"Error uploading to Google Drive: {str(upload_error)}"
                    logger.error(error_msg)
                    logger.exception(upload_error)
                    log_error(error_msg)
                    return {
                        'success': True,
                        'file_path': file_path,
                        'filename': filename,
                        'drive_error': str(upload_error)
                    }
                    
            except Exception as save_error:
                error_msg = f"Error saving document: {str(save_error)}"
                logger.error(error_msg)
                logger.exception(save_error)
                log_error(error_msg)
                return {'success': False, 'error': error_msg}
                
        except Exception as e:
            error_msg = f"Error creating supplemental agreement: {str(e)}"
            logger.error(error_msg)
            logger.exception(e)
            log_error(error_msg)
            return {'success': False, 'error': error_msg}
